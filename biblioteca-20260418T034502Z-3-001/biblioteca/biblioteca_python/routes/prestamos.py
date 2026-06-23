from datetime import datetime, date
from flask import Blueprint, request, jsonify, session, redirect, url_for, render_template, Response
from models.prestamos_model import PrestamosModel
from helpers import str_clean
from security import permission_required

prestamos_bp = Blueprint('prestamos', __name__)


def login_required():
    if not session.get('activo'):
        return redirect(url_for('index'))
    return None


@prestamos_bp.route('/')
def index():
    redir = login_required()
    if redir:
        return redir
    id_user = session['id_usuario']
    m = PrestamosModel()
    perm = m.verificar_permisos(id_user, "Prestamos")
    if not perm and id_user != 1:
        return render_template('permisos.html')
    return render_template('prestamos/index.html')


@prestamos_bp.route('/listar')
@permission_required('Prestamos')
def listar():
    redir = login_required()
    if redir:
        return redir
    m = PrestamosModel()
    data = m.get_prestamos()
    for row in data:
        row['estadoLabel'] = 'Prestado' if row['estado'] == 1 else 'Entregado'
        # Convertir fechas a string si son datetime
        for campo in ('fecha_prestamo', 'fecha_devolucion'):
            if row.get(campo) and not isinstance(row[campo], str):
                row[campo] = str(row[campo])
    return jsonify(data)


@prestamos_bp.route('/registrar', methods=['POST'])
@permission_required('Prestamos')
def registrar():
    redir = login_required()
    if redir:
        return redir
    estudiante = str_clean(request.form.get('estudiante', ''))
    libro = str_clean(request.form.get('libro', ''))
    cantidad = str_clean(request.form.get('cantidad', ''))
    fecha_prestamo = str_clean(request.form.get('fecha_prestamo', ''))
    fecha_devolucion = str_clean(request.form.get('fecha_devolucion', ''))
    observacion = str_clean(request.form.get('observacion', ''))
    if not estudiante or not libro or not cantidad or not fecha_prestamo or not fecha_devolucion:
        return jsonify({'msg': 'Todos los campos son requeridos', 'icono': 'warning'})
    m = PrestamosModel()
    data = m.insertar_prestamo(estudiante, libro, cantidad,
                               fecha_prestamo, fecha_devolucion, observacion)
    if data == 'existe':
        return jsonify({'msg': 'Este estudiante ya tiene este libro prestado', 'icono': 'warning'})
    if data == 'sin_stock':
        return jsonify({'msg': 'Cantidad insuficiente en stock', 'icono': 'warning'})
    if data == 'cantidad_invalida':
        return jsonify({'msg': 'Cantidad invalida', 'icono': 'warning'})
    if data == 'libro_no_encontrado':
        return jsonify({'msg': 'Libro no encontrado', 'icono': 'warning'})
    if data and data > 0:
        return jsonify({'msg': 'Prestamo registrado', 'icono': 'success', 'id': data})
    return jsonify({'msg': 'Error al registrar', 'icono': 'error'})


@prestamos_bp.route('/entregar/<int:id_>')
@permission_required('Prestamos')
def entregar(id_):
    redir = login_required()
    if redir:
        return redir
    m = PrestamosModel()
    data = m.actualizar_prestamo(0, id_)
    if data == 'ok':
        return jsonify({'msg': 'Libro entregado', 'icono': 'success'})
    return jsonify({'msg': 'Error al registrar entrega', 'icono': 'error'})


@prestamos_bp.route('/ticked/<int:id_>')
@permission_required('Prestamos')
def ticked(id_):
    redir = login_required()
    if redir:
        return redir
    m = PrestamosModel()
    data = m.get_prestamo_libro(id_)
    if data:
        for campo in ('fecha_prestamo', 'fecha_devolucion'):
            if data.get(campo) and not isinstance(data[campo], str):
                data[campo] = str(data[campo])
    return jsonify(data)


@prestamos_bp.route('/pdf')
@permission_required('Reportes')
def pdf():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from fpdf import FPDF
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    today = date.today().isoformat()
    prestamo = config_m.get_verificar_prestamos(today)
    if not prestamo:
        return redirect(url_for('configuracion.vacio'))
    pdf_doc = FPDF('P', 'mm', 'letter')
    pdf_doc.add_page()
    pdf_doc.set_margins(10, 10, 10)
    pdf_doc.set_title('Prestamos')
    pdf_doc.set_font('Arial', 'B', 12)
    pdf_doc.cell(195, 5, datos.get('nombre', ''), 0, 1, 'C')
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.cell(20, 5, 'Telefono: ', 0, 0, 'L')
    pdf_doc.set_font('Arial', '', 10)
    pdf_doc.cell(20, 5, datos.get('telefono', ''), 0, 1, 'L')
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.cell(20, 5, 'Correo: ', 0, 0, 'L')
    pdf_doc.set_font('Arial', '', 10)
    pdf_doc.cell(20, 5, datos.get('correo', ''), 0, 1, 'L')
    pdf_doc.ln()
    pdf_doc.set_font('Arial', 'B', 10)
    pdf_doc.set_fill_color(0, 0, 0)
    pdf_doc.set_text_color(255, 255, 255)
    pdf_doc.cell(196, 5, 'Detalle de Prestamos', 1, 1, 'C', True)
    pdf_doc.set_text_color(0, 0, 0)
    pdf_doc.cell(14, 5, 'N', 1, 0, 'L')
    pdf_doc.cell(50, 5, 'Estudiantes', 1, 0, 'L')
    pdf_doc.cell(87, 5, 'Libros', 1, 0, 'L')
    pdf_doc.cell(30, 5, 'Fecha Prestamo', 1, 0, 'L')
    pdf_doc.cell(15, 5, 'Cant.', 1, 1, 'L')
    pdf_doc.set_font('Arial', '', 10)
    for i, row in enumerate(prestamo, 1):
        pdf_doc.cell(14, 5, str(i), 1, 0, 'L')
        pdf_doc.cell(50, 5, row.get('est_nombre', ''), 1, 0, 'L')
        pdf_doc.cell(87, 5, row.get('titulo', ''), 1, 0, 'L')
        pdf_doc.cell(30, 5, str(row.get('fecha_prestamo', '')), 1, 0, 'L')
        pdf_doc.cell(15, 5, str(row.get('cantidad', '')), 1, 1, 'L')
    pdf_bytes = bytes(pdf_doc.output())
    return Response(pdf_bytes, mimetype='application/pdf',
                    headers={'Content-Disposition': 'inline; filename=prestamos.pdf'})


@prestamos_bp.route('/excel')
@permission_required('Reportes')
def excel():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from io import BytesIO
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    config_m = ConfiguracionModel()
    today = date.today().isoformat()
    prestamo = config_m.get_verificar_prestamos(today) or []

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Préstamos"

    # Cabecera
    header_fill = PatternFill("solid", fgColor="1a3a5c")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    headers = ['N°', 'Estudiante', 'Libro', 'Fecha Préstamo', 'Fecha Devolución', 'Cantidad']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center')

    # Datos
    for i, row in enumerate(prestamo, 1):
        ws.append([
            i,
            row.get('est_nombre', ''),
            row.get('titulo', ''),
            str(row.get('fecha_prestamo', '')),
            str(row.get('fecha_devolucion', '')),
            row.get('cantidad', '')
        ])

    # Ancho de columnas
    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename=prestamos_{today}.xlsx'}
    )


@prestamos_bp.route('/word')
@permission_required('Reportes')
def word():
    redir = login_required()
    if redir:
        return redir
    from models.configuracion_model import ConfiguracionModel
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    config_m = ConfiguracionModel()
    datos = config_m.select_configuracion()
    today = date.today().isoformat()
    prestamo = config_m.get_verificar_prestamos(today) or []

    doc = Document()
    # Título
    title = doc.add_heading(datos.get('nombre', 'Biblioteca'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f'Reporte de Préstamos Activos — {today}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Tabla
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['N°', 'Estudiante', 'Libro', 'F. Préstamo', 'Cantidad']):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for i, row in enumerate(prestamo, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row.get('est_nombre', '')
        cells[2].text = row.get('titulo', '')
        cells[3].text = str(row.get('fecha_prestamo', ''))
        cells[4].text = str(row.get('cantidad', ''))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return Response(
        output.read(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=prestamos_{today}.docx'}
    )
